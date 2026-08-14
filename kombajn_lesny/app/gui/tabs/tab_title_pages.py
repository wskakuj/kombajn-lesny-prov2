"""
Kombajn Leśny PRO — Mixin: TabTitlePagesMixin
"""

import customtkinter as ctk
from tkinter import messagebox
from pathlib import Path
import threading
import re
import warnings
import traceback
import pandas as pd
import warnings as _w
_w.filterwarnings("ignore", message=".*OLE2 inconsistency.*")
_w.filterwarnings("ignore", message=".*file size.*not.*sector size.*")
_w.filterwarnings("ignore", message=".*SSCS size.*")
import numpy as np
from docx import Document
import win32com.client
import pythoncom

from app.config import (
    is_file_locked,
)

class TabTitlePagesMixin:
    """Mixin dla ModernApp — metody zostały wyciągnięte z oryginalnego guipia.py."""
    pass

    def setup_mietek_title_pages_tab(self, parent):
        parent.grid_columnconfigure(0, weight=1)
        parent.grid_rowconfigure(0, weight=1)

        scroll_frame = ctk.CTkScrollableFrame(parent, fg_color="transparent")
        scroll_frame.grid(row=0, column=0, sticky="nsew", padx=0, pady=0)
        scroll_frame.grid_columnconfigure(0, weight=1)

        font_label = ctk.CTkFont(family="Segoe UI", size=13, weight="bold")
        font_btn = ctk.CTkFont(family="Segoe UI", size=13)
        card = ctk.CTkFrame(
            scroll_frame,
            fg_color="#252526",
            corner_radius=8,
            border_width=1,
            border_color="#333333",
        )
        card.grid(row=0, column=0, padx=20, pady=(15, 15), sticky="new")
        card.grid_columnconfigure(1, weight=1)
        ctk.CTkLabel(
            card, text="Szablon STR_TYT:", font=font_label, text_color="#E0E0E0"
        ).grid(row=0, column=0, padx=15, pady=(15, 8), sticky="w")
        self.mietek_title_template_entry = ctk.CTkEntry(
            card, placeholder_text="Wskaż plik bazowy", height=36
        )
        self.mietek_title_template_entry.grid(
            row=0, column=1, padx=5, pady=(15, 8), sticky="ew"
        )
        ctk.CTkButton(
            card,
            text="Wybierz Plik",
            command=lambda: self.select_file(
                self.mietek_title_template_entry, [("Word", "*.docx")]
            ),
            width=110,
            height=36,
            font=font_btn,
            fg_color="#333333",
            hover_color="#444444",
        ).grid(row=0, column=2, padx=15, pady=(15, 8))
        ctk.CTkLabel(
            card,
            text="Fold. z plikami Word (OPTAX):",
            font=font_label,
            text_color="#E0E0E0",
        ).grid(row=1, column=0, padx=15, pady=8, sticky="w")
        self.mietek_title_word_entry = ctk.CTkEntry(
            card,
            placeholder_text="Wskaż folder, w którym znajdują się pliki OPTAX",
            height=36,
        )
        self.mietek_title_word_entry.grid(row=1, column=1, padx=5, pady=8, sticky="ew")
        ctk.CTkButton(
            card,
            text="Przeglądaj",
            image=self.icon_folder,
            command=lambda: self.select_dir(self.mietek_title_word_entry),
            width=110,
            height=36,
            font=font_btn,
            fg_color="#333333",
            hover_color="#444444",
        ).grid(row=1, column=2, padx=15, pady=8)
        ctk.CTkLabel(
            card, text="Folder zapisu STR_TYT:", font=font_label, text_color="#E0E0E0"
        ).grid(row=2, column=0, padx=15, pady=8, sticky="w")
        self.mietek_title_output_entry = ctk.CTkEntry(
            card, placeholder_text="Wskaż folder docelowy dla nowych stron", height=36
        )
        self.mietek_title_output_entry.grid(
            row=2, column=1, padx=5, pady=8, sticky="ew"
        )
        ctk.CTkButton(
            card,
            text="Przeglądaj",
            image=self.icon_folder,
            command=lambda: self.select_dir(self.mietek_title_output_entry),
            width=110,
            height=36,
            font=font_btn,
            fg_color="#333333",
            hover_color="#444444",
        ).grid(row=2, column=2, padx=15, pady=8)
        # Zmienne wczytywane w tle - ukryte w GUI
        self.mietek_title_village_placeholder_entry = ctk.CTkEntry(card)
        self.mietek_title_village_placeholder_entry.insert(0, "NAZWA WSI")

        self.mietek_title_area_placeholder_entry = ctk.CTkEntry(card)
        self.mietek_title_area_placeholder_entry.insert(0, "wielkość")
        self.mietek_title_generate_btn = ctk.CTkButton(
            scroll_frame,
            text="Masowo twórz strony STR_TYT",
            image=self.icon_start,
            font=ctk.CTkFont(family="Segoe UI", size=15, weight="bold"),
            fg_color="#0067C0",
            hover_color="#005A9E",
            height=44,
            corner_radius=6,
            command=self.start_mietek_title_pages_pipeline,
        )
        self.mietek_title_generate_btn.grid(
            row=1, column=0, padx=20, pady=(5, 20), sticky="ew"
        )

    def setup_title_pages_tab(self, parent):
        parent.grid_columnconfigure(0, weight=1)
        parent.grid_rowconfigure(0, weight=1)

        scroll_frame = ctk.CTkScrollableFrame(parent, fg_color="transparent")
        scroll_frame.grid(row=0, column=0, sticky="nsew", padx=0, pady=0)
        scroll_frame.grid_columnconfigure(0, weight=1)

        font_label = ctk.CTkFont(family="Segoe UI", size=13, weight="bold")
        font_btn = ctk.CTkFont(family="Segoe UI", size=13)
        card = ctk.CTkFrame(
            scroll_frame,
            fg_color="#252526",
            corner_radius=8,
            border_width=1,
            border_color="#333333",
        )
        card.grid(row=0, column=0, padx=20, pady=(15, 15), sticky="new")
        card.grid_columnconfigure(1, weight=1)
        ctk.CTkLabel(
            card, text="Szablon STR_TYT:", font=font_label, text_color="#E0E0E0"
        ).grid(row=0, column=0, padx=15, pady=(15, 8), sticky="w")
        self.title_template_entry = ctk.CTkEntry(
            card,
            placeholder_text="Wskaż plik bazowy (np. wygenerowany w Kreatorze Szablonów)",
            height=36,
        )
        self.title_template_entry.grid(
            row=0, column=1, padx=5, pady=(15, 8), sticky="ew"
        )
        ctk.CTkButton(
            card,
            text="Wybierz Plik",
            command=lambda: self.select_file(
                self.title_template_entry, [("Word", "*.docx")]
            ),
            width=110,
            height=36,
            font=font_btn,
            fg_color="#333333",
            hover_color="#444444",
        ).grid(row=0, column=2, padx=15, pady=(15, 8))
        ctk.CTkLabel(
            card,
            text="Fold. z rejestrami Excel:",
            font=font_label,
            text_color="#E0E0E0",
        ).grid(row=1, column=0, padx=15, pady=8, sticky="w")
        self.title_excel_entry = ctk.CTkEntry(
            card, placeholder_text="Wskaż folder z plikami .xls / .xlsx", height=36
        )
        self.title_excel_entry.grid(row=1, column=1, padx=5, pady=8, sticky="ew")
        ctk.CTkButton(
            card,
            text="Przeglądaj",
            image=self.icon_folder,
            command=lambda: self.select_dir(self.title_excel_entry),
            width=110,
            height=36,
            font=font_btn,
            fg_color="#333333",
            hover_color="#444444",
        ).grid(row=1, column=2, padx=15, pady=8)
        ctk.CTkLabel(
            card, text="Folder zapisu STR_TYT:", font=font_label, text_color="#E0E0E0"
        ).grid(row=2, column=0, padx=15, pady=8, sticky="w")
        self.title_output_entry = ctk.CTkEntry(
            card, placeholder_text="Wskaż folder docelowy dla nowych stron", height=36
        )
        self.title_output_entry.grid(row=2, column=1, padx=5, pady=8, sticky="ew")
        ctk.CTkButton(
            card,
            text="Przeglądaj",
            image=self.icon_folder,
            command=lambda: self.select_dir(self.title_output_entry),
            width=110,
            height=36,
            font=font_btn,
            fg_color="#333333",
            hover_color="#444444",
        ).grid(row=2, column=2, padx=15, pady=8)
        # Zmienne wczytywane w tle - ukryte w GUI
        self.title_village_placeholder_entry = ctk.CTkEntry(card)
        self.title_village_placeholder_entry.insert(0, "NAZWA WSI")

        self.title_area_placeholder_entry = ctk.CTkEntry(card)
        self.title_area_placeholder_entry.insert(0, "wielkość")
        self.title_generate_btn = ctk.CTkButton(
            scroll_frame,
            text="Masowo twórz strony STR_TYT",
            image=self.icon_start,
            font=ctk.CTkFont(family="Segoe UI", size=15, weight="bold"),
            fg_color="#0067C0",
            hover_color="#005A9E",
            height=44,
            corner_radius=6,
            command=self.start_title_pages_pipeline,
        )
        self.title_generate_btn.grid(row=1, column=0, padx=20, pady=(5, 20), sticky="ew")

    def start_mietek_title_pages_pipeline(self):
        template_path = (
            self.mietek_title_template_entry.get().strip()
            if self.mietek_title_template_entry
            else ""
        )
        word_folder = (
            self.mietek_title_word_entry.get().strip()
            if self.mietek_title_word_entry
            else ""
        )
        output_folder = (
            self.mietek_title_output_entry.get().strip()
            if self.mietek_title_output_entry
            else ""
        )
        village_placeholder = (
            self.mietek_title_village_placeholder_entry.get().strip()
            if self.mietek_title_village_placeholder_entry
            else ""
        )
        area_placeholder = (
            self.mietek_title_area_placeholder_entry.get().strip()
            if self.mietek_title_area_placeholder_entry
            else ""
        )
        if not template_path or not Path(template_path).exists():
            messagebox.showwarning(
                "Błąd", "Wybierz istniejący bazowy plik STR_TYT .docx."
            )
            return
        if not word_folder or not Path(word_folder).exists():
            messagebox.showwarning("Błąd", "Wybierz folder z plikami Word (OPTAX).")
            return
        if not output_folder:
            return
        if not village_placeholder or not area_placeholder:
            return
        if self.running:
            return
        self.last_output_dir = Path(output_folder)
        self._disable_ui_for_process()
        self.log(
            f"[STR_TYT MIETEK] Generowanie stron tytułowych na pods. plików OPTAX..."
        )
        self.set_progress(0)
        threading.Thread(
            target=self.run_mietek_title_pages_thread,
            args=(
                template_path,
                word_folder,
                output_folder,
                village_placeholder,
                area_placeholder,
            ),
            daemon=True,
        ).start()

    def start_title_pages_pipeline(self):
        template_path = (
            self.title_template_entry.get().strip() if self.title_template_entry else ""
        )
        excel_folder = (
            self.title_excel_entry.get().strip() if self.title_excel_entry else ""
        )
        output_folder = (
            self.title_output_entry.get().strip() if self.title_output_entry else ""
        )
        village_placeholder = (
            self.title_village_placeholder_entry.get().strip()
            if self.title_village_placeholder_entry
            else ""
        )
        area_placeholder = (
            self.title_area_placeholder_entry.get().strip()
            if self.title_area_placeholder_entry
            else ""
        )
        if not template_path or not Path(template_path).exists():
            messagebox.showwarning(
                "Błąd", "Wybierz istniejący bazowy plik STR_TYT .docx."
            )
            return
        if not excel_folder or not Path(excel_folder).exists():
            messagebox.showwarning("Błąd", "Wybierz folder z rejestrami Excel.")
            return
        if not output_folder:
            return
        if not village_placeholder or not area_placeholder:
            return
        if self.running:
            return
        self.last_output_dir = Path(output_folder)
        self._disable_ui_for_process()
        self.log(f"[STR_TYT] Tworzenie stron z szablonu: {template_path}")
        self.set_progress(0)
        threading.Thread(
            target=self.run_title_pages_thread,
            args=(
                template_path,
                excel_folder,
                output_folder,
                village_placeholder,
                area_placeholder,
            ),
            daemon=True,
        ).start()

    def run_mietek_title_pages_thread(
            self,
            template_path_str,
            word_folder_str,
            output_folder_str,
            village_placeholder,
            area_placeholder,
    ):
        pythoncom.CoInitialize()
        word = None
        try:
            template_path = Path(template_path_str)
            word_folder = Path(word_folder_str)
            output_folder = Path(output_folder_str)
            output_folder.mkdir(parents=True, exist_ok=True)
            files = sorted(
                [
                    p
                    for p in word_folder.rglob("OPTAX*.doc*")
                    if p.is_file() and not p.name.startswith("~$")
                ]
            )
            total = len(files)
            created = 0
            if total == 0:
                self.log("Brak plików 'OPTAX*.doc*' we wskazanym folderze.")
            else:
                word = win32com.client.DispatchEx("Word.Application")
                word.Visible = False
                word.DisplayAlerts = 0
                self.start_progress_tracking(total, "Generowanie STR_TYT (MIETEK)")

                for idx, file_path in enumerate(files, start=1):
                    self.check_stop()
                    self.progress_current_file = file_path.name
                    try:
                        if is_file_locked(file_path):
                            self.log(f"Pominięto (zablokowany): {file_path.name}")
                            continue
                        doc_word = word.Documents.Open(str(file_path), ReadOnly=True)
                        text_content = doc_word.Content.Text
                        doc_word.Close(SaveChanges=False)
                        village_match = re.search(
                            r"Obiekt:\s*(.+?)(?=\s{2,}|\t|\r|\n|$)",
                            text_content,
                            re.IGNORECASE,
                        )
                        area_match = re.search(
                            r"Razem\s*[^0-9A-Za-z]*([\d\s]+(?:[\.,]\d+)?)",
                            text_content,
                            re.IGNORECASE,
                        )
                        if village_match:
                            village_name = village_match.group(1).strip().upper()
                        else:
                            self.log(
                                f"UWAGA: Nie odnaleziono nazwy wsi (Obiekt:) w pliku: {file_path.name}"
                            )
                            village_name = "NIEZNANA_WIES"
                        if area_match:
                            area_str = area_match.group(1).replace(" ", "").strip()
                        else:
                            self.log(
                                f"UWAGA: Nie odnaleziono pola powierzchni (Razem │) w pliku: {file_path.name}"
                            )
                            area_str = "[BRAK_DANYCH]"
                        doc = Document(str(template_path))
                        self.replace_text_robust(doc, village_placeholder, village_name)
                        self.replace_text_robust(doc, area_placeholder, area_str)
                        safe_village_name = "".join(
                            [
                                c
                                for c in village_name
                                if c.isalpha() or c.isdigit() or c in " -_"
                            ]
                        ).strip()
                        new_doc_name = f"STR_TYT_{safe_village_name}.docx"
                        doc.save(str(output_folder / new_doc_name))
                        created += 1
                        self.log(f"Utworzono: {new_doc_name} (Pow: {area_str})")
                    except Exception as e:
                        self.log(f"Błąd podczas obróbki pliku {file_path.name}: {e}")
                    self.set_progress(idx / total)
                self.update_status("Zakończono", "#27ae60", animate=False)
                self.after(
                    0,
                    lambda: messagebox.showinfo(
                        "Sukces",
                        f"Zakończono. Utworzono {created} stron STR_TYT z plików MIETEK (OPTAX).",
                    ),
                )
        except InterruptedError:
            self.update_status("Przerwano", "#D83B01", animate=False)
            self.log("\nZADANIE PRZERWANE PRZEZ UŻYTKOWNIKA.")
        except Exception as e:
            self.log(traceback.format_exc())
            self.update_status("Błąd", "#D83B01", animate=False)
        finally:
            if word is not None:
                try:
                    word.Quit()
                except:
                    pass
            pythoncom.CoUninitialize()
            self.running = False
            self.after(0, self.restore_all_buttons)

    def run_title_pages_thread(
            self,
            template_path_str,
            excel_folder_str,
            output_folder_str,
            village_placeholder,
            area_placeholder,
    ):
        try:
            template_path = Path(template_path_str)
            excel_folder = Path(excel_folder_str)
            output_folder = Path(output_folder_str)
            files = sorted(
                [
                    p
                    for p in excel_folder.iterdir()
                    if p.is_file()
                       and p.suffix.lower() in {".xls", ".xlsx"}
                       and not p.name.startswith("~$")
                ]
            )
            total = len(files)
            created = 0
            self.start_progress_tracking(total, "Generowanie STR_TYT")

            for idx, file_path in enumerate(files, start=1):
                self.check_stop()
                self.progress_current_file = file_path.name
                try:
                    if is_file_locked(file_path):
                        self.log(f"Pominięto (zablokowany raport): {file_path.name}")
                        continue
                    village_name = self.extract_village_name_from_excel(file_path.name)
                    area_str = self.read_area_from_excel(file_path)
                    if area_str is None:
                        continue
                    doc = Document(str(template_path))
                    self.replace_text_robust(
                        doc, village_placeholder, village_name.upper()
                    )
                    self.replace_text_robust(doc, area_placeholder, area_str)
                    new_doc_name = f"STR_TYT_{village_name}.docx"
                    doc.save(str(output_folder / new_doc_name))
                    created += 1
                except Exception as e:
                    self.log(f"Błąd {file_path.name}: {e}")
                self.set_progress(idx / total)
            self.update_status("Zakończono", "#27ae60", animate=False)
            self.after(
                0,
                lambda: messagebox.showinfo(
                    "Sukces", f"Zakończono. Utworzono {created} stron STR_TYT."
                ),
            )
        except InterruptedError:
            self.update_status("Przerwano", "#D83B01", animate=False)
            self.log("\nZADANIE PRZERWANE PRZEZ UŻYTKOWNIKA.")
        except Exception as e:
            self.log(traceback.format_exc())
            self.update_status("Błąd", "#D83B01", animate=False)
        finally:
            self.running = False
            self.after(0, self.restore_all_buttons)

    def replace_in_paragraph(self, paragraph, old_text, new_text):
        if not old_text:
            return
        if old_text.lower() in paragraph.text.lower():
            full_text = "".join(run.text for run in paragraph.runs)
            pattern = re.compile(re.escape(old_text), re.IGNORECASE)
            new_full_text = pattern.sub(new_text, full_text)
            if paragraph.runs:
                paragraph.runs[0].text = new_full_text
                for i in range(1, len(paragraph.runs)):
                    paragraph.runs[i].text = ""

    def replace_text_robust(self, doc, old_text, new_text):
        for paragraph in doc.paragraphs:
            self.replace_in_paragraph(paragraph, old_text, new_text)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        self.replace_in_paragraph(paragraph, old_text, new_text)
        for section in doc.sections:
            for paragraph in section.header.paragraphs:
                self.replace_in_paragraph(paragraph, old_text, new_text)
            for paragraph in section.footer.paragraphs:
                self.replace_in_paragraph(paragraph, old_text, new_text)

    def extract_village_name_from_excel(self, filename):
        match = re.search(r"-\d+-(.+?)\.(xls|xlsx)$", filename, re.IGNORECASE)
        if match:
            return match.group(1)
        name_without_ext = Path(filename).stem
        parts = name_without_ext.split("-")
        return parts[-1] if len(parts) > 1 else name_without_ext

    def read_area_from_excel(self, excel_file_path):
        df = pd.read_excel(excel_file_path, sheet_name="OT", header=None)
        first_col = df.iloc[:, 0].astype(str).str.strip().str.lower()
        row_idx = df[first_col == "ogółem:"].index
        if row_idx.empty:
            return None
        area_val = df.iloc[row_idx[0], 1]
        if pd.isna(area_val):
            return "[brak danych]"
        if isinstance(area_val, float):
            area_val = round(area_val, 4)
            return str(int(area_val)) if area_val.is_integer() else str(area_val)
        return str(area_val)

