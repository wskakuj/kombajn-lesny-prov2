"""
Kombajn Leśny PRO — Mixin: TabTemplateGeneratorMixin
"""

import customtkinter as ctk
import tkinter as tk
from tkinter import messagebox
from pathlib import Path
import numpy as np
from docx import Document

from app.config import (
    TERRITORY_DATA,
)

from app.core.word_worker import (
    get_resource_path, replace_text_in_tables, replace_text_preserve_runs,
)

class TabTemplateGeneratorMixin:
    """Mixin dla ModernApp — metody zostały wyciągnięte z oryginalnego guipia.py."""
    pass

    def setup_template_generator_tab(self, parent, mode_key):
        scroll_frame = ctk.CTkScrollableFrame(parent, fg_color="transparent")
        scroll_frame.pack(fill="both", expand=True, padx=10, pady=10)
        scroll_frame.grid_columnconfigure(0, weight=1)
        font_label = ctk.CTkFont(family="Segoe UI", size=13, weight="bold")
        card = ctk.CTkFrame(
            scroll_frame,
            fg_color="#252526",
            corner_radius=8,
            border_width=1,
            border_color="#333333",
        )
        card.grid(row=0, column=0, sticky="new", pady=(0, 10))
        card.grid_columnconfigure((1, 3), weight=1)

        ctk.CTkLabel(card, text="Typ dokumentu:", font=font_label).grid(
            row=0, column=0, padx=15, pady=(15, 8), sticky="e"
        )
        self.tpl_data[mode_key]["doc_type_var"] = ctk.StringVar(value="UPUL")
        opt_doc = ctk.CTkOptionMenu(
            card,
            variable=self.tpl_data[mode_key]["doc_type_var"],
            values=["UPUL", "ISL"],
            width=200,
            height=32,
        )
        opt_doc.grid(row=0, column=1, padx=10, pady=(15, 8), sticky="w")

        ctk.CTkLabel(card, text="Prefiks obrębu:", font=font_label).grid(
            row=0, column=2, padx=15, pady=(15, 8), sticky="e"
        )
        self.tpl_data[mode_key]["prefix_var"] = ctk.StringVar(
            value="położonych na terenie obrębu"
        )
        opt_pref = ctk.CTkOptionMenu(
            card,
            variable=self.tpl_data[mode_key]["prefix_var"],
            values=["położonych na terenie obrębu", "Obręb:"],
            width=230,
            height=32,
        )
        opt_pref.grid(row=0, column=3, padx=15, pady=(15, 8), sticky="w")

        # --- KASKADOWE LISTY TERYTORIALNE ---
        ctk.CTkLabel(card, text="Województwo:", font=font_label).grid(
            row=1, column=0, padx=15, pady=8, sticky="e"
        )
        woj_list = sorted(TERRITORY_DATA.keys())
        if not woj_list:
            woj_list = ["BRAK DANYCH"]

        default_woj = (
            "KUJAWSKO-POMORSKIE"
            if "KUJAWSKO-POMORSKIE" in TERRITORY_DATA
            else woj_list[0]
        )

        self.tpl_data[mode_key]["woj_var"] = ctk.StringVar(value=default_woj)
        self.tpl_data[mode_key]["woj_menu"] = ctk.CTkOptionMenu(
            card,
            variable=self.tpl_data[mode_key]["woj_var"],
            values=woj_list,
            width=230,
            height=32,
            command=lambda val, mk=mode_key: self._on_woj_change(mk, val),
        )
        self.tpl_data[mode_key]["woj_menu"].grid(
            row=1, column=1, padx=10, pady=8, sticky="ew"
        )

        ctk.CTkLabel(card, text="Powiat:", font=font_label).grid(
            row=1, column=2, padx=15, pady=8, sticky="e"
        )
        powiat_list = sorted(TERRITORY_DATA.get(default_woj, {}).keys())
        if not powiat_list:
            powiat_list = ["BRAK DANYCH"]

        default_powiat = (
            "TUCHOLSKI"
            if "TUCHOLSKI" in TERRITORY_DATA.get(default_woj, {})
            else powiat_list[0]
        )

        self.tpl_data[mode_key]["powiat_var"] = ctk.StringVar(value=default_powiat)
        self.tpl_data[mode_key]["powiat_menu"] = ctk.CTkComboBox(
            card,
            variable=self.tpl_data[mode_key]["powiat_var"],
            values=powiat_list,
            width=230,
            height=32,
            command=lambda val, mk=mode_key: self._on_powiat_change(mk, val),
        )
        self.tpl_data[mode_key]["powiat_menu"].grid(
            row=1, column=3, padx=15, pady=8, sticky="ew"
        )

        ctk.CTkLabel(card, text="Gmina:", font=font_label).grid(
            row=2, column=0, padx=15, pady=8, sticky="e"
        )
        gmina_list = TERRITORY_DATA.get(default_woj, {}).get(default_powiat, [])
        if not gmina_list:
            gmina_list = ["BRAK DANYCH"]

        default_gmina = "LUBIEWO" if "LUBIEWO" in gmina_list else gmina_list[0]
        self.tpl_data[mode_key]["gmina_var"] = ctk.StringVar(value=default_gmina)
        self.tpl_data[mode_key]["gmina_menu"] = ctk.CTkComboBox(
            card,
            variable=self.tpl_data[mode_key]["gmina_var"],
            values=gmina_list,
            width=230,
            height=32,
        )
        self.tpl_data[mode_key]["gmina_menu"].grid(
            row=2, column=1, padx=10, pady=8, sticky="ew"
        )
        # ------------------------------------

        ctk.CTkLabel(card, text="Stan na:", font=font_label).grid(
            row=2, column=2, padx=15, pady=8, sticky="e"
        )
        self.tpl_data[mode_key]["stan_na_entry"] = ctk.CTkEntry(card, height=32)
        self.tpl_data[mode_key]["stan_na_entry"].insert(0, "30.06.2026 r.")
        self.tpl_data[mode_key]["stan_na_entry"].grid(
            row=2, column=3, padx=15, pady=8, sticky="ew"
        )

        ctk.CTkLabel(card, text="Na okres:", font=font_label).grid(
            row=3, column=0, padx=15, pady=8, sticky="e"
        )
        self.tpl_data[mode_key]["okres_entry"] = ctk.CTkEntry(card, height=32)
        self.tpl_data[mode_key]["okres_entry"].insert(0, "01.01.2027 – 31.12.2036 r.")
        self.tpl_data[mode_key]["okres_entry"].grid(
            row=3, column=1, padx=10, pady=8, sticky="ew"
        )

        self.tpl_data[mode_key]["single_village_var"] = ctk.BooleanVar(value=False)
        cb_single = ctk.CTkCheckBox(
            card,
            text="Stwórz stronę dla konkretnej wsi",
            variable=self.tpl_data[mode_key]["single_village_var"],
            command=lambda: self._toggle_single_village(mode_key)
        )
        cb_single.grid(row=3, column=2, columnspan=2, padx=15, pady=8, sticky="w")

        self.tpl_data[mode_key]["area_var"] = ctk.BooleanVar(value=True)
        cb_area = ctk.CTkCheckBox(
            card,
            text="Dodaj wiersz z powierzchnią (ha)",
            variable=self.tpl_data[mode_key]["area_var"],
            command=lambda: self._sync_area_row_state(mode_key),
        )
        cb_area.grid(row=4, column=0, columnspan=2, padx=15, pady=8, sticky="w")

        self.tpl_data[mode_key]["lbl_village"] = ctk.CTkLabel(card, text="Nazwa wsi:", font=font_label)
        self.tpl_data[mode_key]["village_entry"] = ctk.CTkEntry(card, height=32)
        self.tpl_data[mode_key]["village_entry"].insert(0, "NAZWA WSI")

        self.tpl_data[mode_key]["lbl_village"].grid(row=5, column=0, padx=15, pady=8, sticky="e")
        self.tpl_data[mode_key]["village_entry"].grid(row=5, column=1, padx=10, pady=8, sticky="ew")

        self.tpl_data[mode_key]["lbl_area"] = ctk.CTkLabel(card, text="Powierzchnia:", font=font_label)
        self.tpl_data[mode_key]["area_entry"] = ctk.CTkEntry(card, height=32)
        self.tpl_data[mode_key]["area_entry"].insert(0, "wielkość")
        self.tpl_data[mode_key]["area_entry"].bind("<KeyRelease>", lambda e: self._sync_area_row_state(mode_key))

        self.tpl_data[mode_key]["lbl_area"].grid(row=5, column=2, padx=15, pady=8, sticky="e")
        self.tpl_data[mode_key]["area_entry"].grid(row=5, column=3, padx=15, pady=8, sticky="ew")

        ctk.CTkLabel(
            card, text="Zapisz szablon jako:", font=font_label, text_color="#E0E0E0"
        ).grid(row=6, column=0, padx=15, pady=(15, 20), sticky="e")
        self.tpl_data[mode_key]["output_entry"] = ctk.CTkEntry(
            card, placeholder_text="Ścieżka do pliku np. Szablon.docx", height=32
        )
        self.tpl_data[mode_key]["output_entry"].grid(
            row=6, column=1, columnspan=2, padx=10, pady=(15, 20), sticky="ew"
        )
        btn_browse = ctk.CTkButton(
            card,
            text="Wybierz...",
            image=self.icon_folder,
            width=110,
            height=32,
            fg_color="#333333",
            hover_color="#444444",
            command=lambda m=mode_key: self.select_save_file(
                self.tpl_data[m]["output_entry"]
            ),
        )
        btn_browse.grid(row=6, column=3, padx=15, pady=(15, 20), sticky="w")

        self.tpl_data[mode_key]["btn_gen"] = ctk.CTkButton(
            scroll_frame,
            text="Wygeneruj Szablon STR_TYT",
            image=self.icon_start,
            font=ctk.CTkFont(size=15, weight="bold"),
            fg_color="#27ae60",
            hover_color="#219653",
            height=44,
            corner_radius=6,
            command=lambda m=mode_key: self.generate_template_now(m),
        )
        self.tpl_data[mode_key]["btn_gen"].grid(
            row=1, column=0, pady=(5, 10), sticky="ew"
        )

        # Wywołujemy ukrycie pól na start
        self._toggle_single_village(mode_key)
        self._sync_area_row_state(mode_key)

    def generate_template_now(self, mode_key):
        vars_dict = self.tpl_data[mode_key]
        doc_type = vars_dict["doc_type_var"].get()
        prefix = vars_dict["prefix_var"].get().strip()
        gmina = vars_dict["gmina_var"].get().strip().upper()
        powiat = vars_dict["powiat_var"].get().strip().upper()
        woj = vars_dict["woj_var"].get().strip().upper()
        stan_na = vars_dict["stan_na_entry"].get().strip()
        okres = vars_dict["okres_entry"].get().strip()

        # --- NOWA LOGIKA DLA WSI I POWIERZCHNI ---
        if vars_dict.get("single_village_var") and vars_dict["single_village_var"].get():
            village = vars_dict["village_entry"].get().strip().upper() or "NAZWA WSI"
            area_text = vars_dict["area_entry"].get().strip() if vars_dict["area_var"].get() else ""
        else:
            village = "NAZWA WSI"
            area_text = "wielkość" if vars_dict["area_var"].get() else ""
        # ------------------------------------------

        out_path = vars_dict["output_entry"].get().strip()
        if not out_path:
            messagebox.showwarning(
                "Błąd", "Wskaż miejsce i nazwę pliku do zapisu (np. Mojszablon.docx)!"
            )
            return
        if vars_dict["area_var"].get() and not area_text:
            messagebox.showwarning(
                "Błąd",
                "Wpisz wartość dla pola Powierzchnia albo odznacz 'Dodaj wiersz z powierzchnią (ha)'.",
            )
            return
        # Nazwa pliku zawsze z przedrostkiem zależnym od typu dokumentu (UPUL_ / ISL_)
        doc_prefix = "ISL_" if doc_type == "ISL" else "UPUL_"
        other_prefix = "UPUL_" if doc_type == "ISL" else "ISL_"
        out_path_obj = Path(out_path)
        file_name = out_path_obj.name
        if file_name.upper().startswith(other_prefix):
            file_name = file_name[len(other_prefix):]
        if not file_name.upper().startswith(doc_prefix):
            file_name = f"{doc_prefix}{file_name}"
        if not file_name.lower().endswith(".docx"):
            file_name = f"{file_name}.docx"
        out_path = str(out_path_obj.with_name(file_name))
        vars_dict["output_entry"].delete(0, "end")
        vars_dict["output_entry"].insert(0, out_path)
        sample_name = "STR_TYT.docx"  # Zmienione na Twoją nową nazwę pliku bazowego
        sample_path = get_resource_path(sample_name)
        if not sample_path.exists():
            messagebox.showerror(
                "Brak wzorca",
                f"Nie znaleziono wbudowanego pliku wzorcowego: {sample_name}",
            )
            return
        try:
            doc = Document(sample_path)
            replacements = {
                "LUBIEWO": gmina,
                "TUCHOLSKI": powiat,
                "KUJAWSKO-POMORSKIE": woj,
                "30.06.2026 r.": stan_na,
                "01.01.2027 – 31.12.2036 r.": okres,
                "NAZWA WSI": village,
            }
            if doc_type == "ISL":
                replacements["UPROSZCZONY PLAN URZĄDZANIA LASÓW"] = (
                    "INWENTARYZACJA STANU LASU"
                )
                replacements["nie stanowiących własności Skarbu Państwa"] = (
                    "dla lasów niestanowiących własności Skarbu Państwa"
                )
            if prefix == "Obręb:":
                replacements["położonych na terenie"] = ""
                replacements["obrębu"] = "Obręb:"
            for paragraph in doc.paragraphs:
                replace_text_preserve_runs(paragraph, replacements)
            replace_text_in_tables(doc.tables, replacements)
            try:
                self._apply_area_toggle_to_doc(doc, mode_key)
            except Exception as e:
                self.log(
                    f"Ostrzeżenie: Błąd podczas formatowania pola powierzchni: {e}"
                )
            doc.save(out_path)
            self.last_output_dir = Path(out_path).parent
            self.open_dir_btn.configure(state="normal")
            self.log(
                f"[KREATOR SZABLONU] Zapisano nowy szablon bazowy na podstawie wzorca: {out_path}"
            )
            messagebox.showinfo(
                "Sukces",
                "Szablon wygenerowany pomyślnie. Zachowano układ, czcionki i logo.",
            )
        except Exception as e:
            self.log(f"Błąd podczas tworzenia szablonu: {e}")
            messagebox.showerror("Błąd", f"Nie udało się wygenerować szablonu:\n{e}")

    def _apply_area_toggle_to_doc(self, doc, mode_key):
        vars_dict = self.tpl_data[mode_key]
        keep_area = bool(vars_dict["area_var"].get())
        area_text = (
            vars_dict["area_entry"].get().strip() if "area_entry" in vars_dict else ""
        )
        keywords = ["ogólna opracowania"]
        if keep_area:
            if not area_text:
                return
            for p in doc.paragraphs:
                self.replace_in_paragraph(p, "powierzchnia", area_text)
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for p in cell.paragraphs:
                            self.replace_in_paragraph(p, "powierzchnia", area_text)
        else:
            for p in list(doc.paragraphs):
                if any(kw.lower() in p.text.lower() for kw in keywords):
                    p._element.getparent().remove(p._element)
            for table in doc.tables:
                for row in list(table.rows):
                    row_text = " ".join(c.text for c in row.cells).lower()
                    if any(kw.lower() in row_text for kw in keywords):
                        tr = row._tr
                        tr.getparent().remove(tr)
        if keep_area:
            if not area_text:
                return
            replacements = {"Powierzchnia": area_text, "powierzchnia": area_text}
            for p in doc.paragraphs:
                replace_text_preserve_runs(p, replacements)
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for p in cell.paragraphs:
                            replace_text_preserve_runs(p, replacements)
        else:
            for p in list(doc.paragraphs):
                if any(kw in p.text for kw in keywords):
                    p._element.getparent().remove(p._element)
            for table in doc.tables:
                for row in list(table.rows):
                    row_text = " ".join(c.text for c in row.cells)
                    if any(kw in row_text for kw in keywords):
                        tr = row._tr
                        tr.getparent().remove(tr)

    def _toggle_single_village(self, mode_key):
        vars_dict = self.tpl_data[mode_key]
        if vars_dict["single_village_var"].get():
            vars_dict["lbl_village"].grid()
            vars_dict["village_entry"].grid()
            if vars_dict["area_var"].get():
                vars_dict["lbl_area"].grid()
                vars_dict["area_entry"].grid()
        else:
            vars_dict["lbl_village"].grid_remove()
            vars_dict["village_entry"].grid_remove()
            vars_dict["lbl_area"].grid_remove()
            vars_dict["area_entry"].grid_remove()

    def _sync_area_row_state(self, mode_key):
        vars_dict = self.tpl_data[mode_key]
        if "area_entry" not in vars_dict:
            return

        state = "normal" if vars_dict["area_var"].get() else "disabled"
        try:
            vars_dict["area_entry"].configure(state=state)
        except Exception:
            pass

        if not vars_dict["area_var"].get():
            try:
                vars_dict["area_entry"].delete(0, tk.END)
            except Exception:
                pass

        if vars_dict.get("single_village_var") and vars_dict["single_village_var"].get():
            if vars_dict["area_var"].get():
                vars_dict["lbl_area"].grid()
                vars_dict["area_entry"].grid()
            else:
                vars_dict["lbl_area"].grid_remove()
                vars_dict["area_entry"].grid_remove()

    def _toggle_global_font_size(self):
        """Włącza/wyłącza globalne ustawienie czcionki i zarządza stanem pól"""
        if self.global_font_var.get():
            # Zaznaczony - włącz globalne pole, wyłącz wszystkie indywidualne
            self.global_font_entry.configure(
                state="normal",
                fg_color="#252526",
                text_color="#FFFFFF",
                border_color="#0078D7",
            )
            for sheet_name, data in self.excel_font_entries.items():
                data["entry"].configure(
                    state="disabled",
                    fg_color="#2A2A2A",
                    text_color="#666666",
                    border_color="#333333",
                )
        else:
            # Odznaczony - wyłącz globalne pole, włącz wszystkie indywidualne
            self.global_font_entry.configure(
                state="disabled",
                fg_color="#2A2A2A",
                text_color="#666666",
                border_color="#333333",
            )
            for sheet_name, data in self.excel_font_entries.items():
                data["entry"].configure(
                    state="normal",
                    fg_color="#252526",
                    text_color="#FFFFFF",
                    border_color="#333333",
                )

    def _on_woj_change(self, mode_key, selected_woj):
        """Aktualizuje listę powiatów po zmianie województwa"""
        powiaty = sorted(TERRITORY_DATA.get(selected_woj, {}).keys())
        if not powiaty:
            powiaty = ["Brak powiatów"]
        self.tpl_data[mode_key]["powiat_menu"].configure(values=powiaty)
        self.tpl_data[mode_key]["powiat_var"].set(powiaty[0])
        self._on_powiat_change(mode_key, powiaty[0])

    def _on_powiat_change(self, mode_key, selected_powiat):
        """Aktualizuje listę gmin po zmianie powiatu"""
        selected_woj = self.tpl_data[mode_key]["woj_var"].get()
        gminy = TERRITORY_DATA.get(selected_woj, {}).get(selected_powiat, ["Brak gmin"])
        self.tpl_data[mode_key]["gmina_menu"].configure(values=gminy)
        self.tpl_data[mode_key]["gmina_var"].set(gminy[0])

