"""
Kombajn Leśny PRO — Proces wykonawczy Worda (COM)
===================================================
Zależności: config.py (ENCODING, MACRO_MAP, SEQUENCES_TO_REMOVE, clean_xml_incompatible, COLORS)
Odpowiada za: generowanie plików DOC/DOCX z TXT, aplikowanie makr VBA,
              przetwarzanie REJESTR1, konwersję przez Microsoft Word COM.

Funkcja run_word_worker jest projektowana do uruchamiania w osobnym procesie
(lub wątku), aby nie blokować GUI.
"""

import os
import re
import sys
import time
import json
import subprocess
from pathlib import Path

from docx import Document
from docx.shared import Pt
from docx.shared import Cm
from docx.oxml.ns import qn
import win32com.client

from app.config import (
    ENCODING, MACRO_MAP, SEQUENCES_TO_REMOVE,
    clean_xml_incompatible, COLORS,
)

def run_word_worker(in_dir_str, out_dir_str, remove_names, file_filter=None, margin_config=None):
    if margin_config is None: margin_config = {}

    if sys.stdout is not None and hasattr(sys.stdout, "reconfigure"):
        sys.stdout.reconfigure(encoding="utf-8", errors="replace")
    if sys.stderr is not None and hasattr(sys.stderr, "reconfigure"):
        sys.stderr.reconfigure(encoding="utf-8", errors="replace")

    import pyautogui

    FILES_TO_FIX = ["OPTAX", "TAB_KLW3", "HALIZNY", "WYK_NEG"]
    in_dir = Path(in_dir_str)
    out_dir = Path(out_dir_str)
    files = list(in_dir.rglob("*.txt"))

    selected_filters = []
    if file_filter:
        if isinstance(file_filter, (list, tuple, set)):
            selected_filters = [str(x).upper() for x in file_filter if str(x).strip()]
        else:
            selected_filters = [str(file_filter).upper()]

    if selected_filters and "WSZYSTKIE" not in selected_filters:
        files = [f for f in files if f.stem.upper() in selected_filters]

    if not files:
        print(f"[INFO] Brak plików TXT do przetworzenia dla filtru: {file_filter}")
        return

    print(">>> Generowanie plików DOCX (REJESTR)...")
    for f in files:
        if f.stem.upper() == "REJESTR1":
            rel_path = f.relative_to(in_dir)
            target = out_dir / rel_path.parent / f"{f.stem}.docx"
            target.parent.mkdir(parents=True, exist_ok=True)
            with open(f, "r", encoding=ENCODING, errors="ignore") as fh:
                text = fh.read()
            if "\x00" in text:
                continue

            # --- ZAMIAST POWOLNEGO VBA: BŁYSKAWICZNE PRZETWARZANIE W PYTHONIE ---
            if remove_names:
                # 1. Odpowiednik "ZamienLF" (pionowy tab -> nowa linia + twarda spacja)
                text = text.replace("\x0b", "\n\xa0")

                # 2. Odpowiednik wycinania kolumny od dołu (w Pythonie tniemy po linijkach)
                lines = text.split("\n")
                for i in range(2, len(lines)):
                    if len(lines[i]) >= 62:  # 11 (StartUsun) + 51 (DlugoscUsun) = 62
                        # Zostawiamy 10 pierwszych znaków i łączymy z resztą tekstu za nazwiskiem
                        lines[i] = lines[i][:10] + lines[i][61:]
                text = "\n".join(lines)

                # 3. Odpowiedniki "Replace" (dokładnie jak w VBA)
                text = text.replace('AGENCJA "C', 'AGENCJA "CEZAR"')
                text = text.replace('Rejestr deli',
                                    'Rejestr działek leśnych i gruntów do zalesienia wg. właścicieli')
                text = text.replace('AGENCJA „CEZAR"EZAR"', 'AGENCJA "CEZAR"')
                text = text.replace('AGENCJA', 'AGENCJA "CEZAR"')
                text = text.replace('AGENCJA "CEZAR" "CEZAR"', 'AGENCJA "CEZAR"')
                text = text.replace('Rejestr ieli',
                                    'Rejestr działek leśnych i gruntów do zalesienia wg. właścicieli')
                text = text.replace("Wskazania godspodarcze", " Wskazania gospodarcze")
                text = re.sub(r'E\s*$', '', text)
            # --------------------------------------------------------------------

            doc = Document()

            # --- ZASTOSOWANIE MARGINESÓW DLA REJESTR1 ---
            if "REJESTR1" in margin_config:
                m = margin_config["REJESTR1"]
                for section in doc.sections:
                    section.top_margin = Cm(m[0])
                    section.bottom_margin = Cm(m[1])
                    section.left_margin = Cm(m[2])
                    section.right_margin = Cm(m[3])
            # --------------------------------------------

            doc.styles["Normal"].font.name = "Cascadia Code"
            doc.styles["Normal"].font.size = Pt(10)
            doc.styles["Normal"].element.rPr.rFonts.set(qn("w:eastAsia"), "Cascadia Code")
            p = doc.add_paragraph()
            r = p.add_run(clean_xml_incompatible(text))
            r.font.name = "Cascadia Code"
            r.font.size = Pt(10)
            doc.save(str(target))
            print(f"  └─ Utworzono: {target.parent.name}/{target.name}")

    print(">>> Przygotowywanie środowiska Microsoft Word...")
    word = None  # <--- Inicjalizacja przed try
    try:
        word = win32com.client.DispatchEx("Word.Application")
        word.Visible = False  # <--- ZMIANA NA True (Word będzie widoczny)
        word.DisplayAlerts = 0  # <--- ZMIANA NA -1 (Włączamy alerty)
        shell = win32com.client.Dispatch("WScript.Shell")
        print(">>> Generowanie standardowych plików DOC...")
        for f in files:
            if f.stem.upper() != "REJESTR1":
                rel_path = f.relative_to(in_dir)
                target = out_dir / rel_path.parent / f"{f.stem}.doc"
                target.parent.mkdir(parents=True, exist_ok=True)
                with open(f, "r", encoding=ENCODING, errors="replace") as fh:
                    text = fh.read()

                doc = word.Documents.Add()

                # --- ZASTOSOWANIE MARGINESÓW DLA POZOSTAŁYCH ---
                f_upper = f.stem.upper()
                if f_upper in margin_config:
                    m = margin_config[f_upper]
                    doc.PageSetup.TopMargin = 28.35 * m[0]
                    doc.PageSetup.BottomMargin = 28.35 * m[1]
                    doc.PageSetup.LeftMargin = 28.35 * m[2]
                    doc.PageSetup.RightMargin = 28.35 * m[3]
                # -----------------------------------------------

                doc.Content.InsertAfter(text)
                doc.Content.Font.Name = "Cascadia Code"
                doc.Content.Font.Size = 10
                doc.Content.ParagraphFormat.SpaceAfter = 0
                doc.Content.ParagraphFormat.SpaceBefore = 0
                doc.Content.ParagraphFormat.LineSpacingRule = 0
                doc.SaveAs(str(target), FileFormat=0)
                doc.Close(SaveChanges=False)
                print(f"  └─ Utworzono: {target.parent.name}/{target.name}")

        print(">>> Aplikowanie makr dla standardowych plików DOC...")
        for f in files:
            macro_name = MACRO_MAP.get(f.stem.lower())
            if macro_name:
                rel_path = f.relative_to(in_dir)
                target = out_dir / rel_path.parent / f"{f.stem}.doc"
                if target.exists():
                    doc = word.Documents.Open(str(target))
                    try:
                        word.Run(macro_name)
                        print(
                            f"  └─ [{macro_name}] Zastosowano na: {target.parent.name}/{target.name}"
                        )
                    except Exception as e:
                        print(
                            f"  └─ Ostrzeżenie: Makro [{macro_name}] nie zadziałało ({e})"
                        )
                    finally:
                        doc.Close(SaveChanges=True)

        print(">>> Przetwarzanie i układanie REJESTRU...")
        for f in files:
            if f.stem.upper() == "REJESTR1":
                rel_path = f.relative_to(in_dir)
                target = out_dir / rel_path.parent / f"{f.stem}.docx"
                if target.exists():
                    doc = word.Documents.Open(str(target))
                    try:
                        # Wywołujemy już TYLKO przesunięcie (resztę odwalił Python w pamięci)
                        try:
                            word.Run("REJESTR_Z_PRZESUWANIEM")
                            time.sleep(0.5)
                            for _ in range(2):
                                shell.SendKeys("{ENTER}", 0)
                                time.sleep(0.5)
                            if doc.Content.End > 2:
                                doc.Range(Start=0, End=2).Delete()
                        except Exception as e:
                            print(
                                f"  └─ Ostrzeżenie: Błąd makra REJESTR_Z_PRZESUWANIEM ({e})"
                            )

                        if remove_names:
                            try:
                                # Ustawiamy globalnie wyrównanie tekstu, czcionkę 9 oraz parametry akapitu
                                doc.Content.ParagraphFormat.Alignment = 1  # 1 = wdAlignParagraphCenter
                                doc.Content.Font.Size = 9

                                # Opcje łamania i odstępów, z którymi normalnie radził sobie stary makro-kod ZamienLF
                                doc.Content.ParagraphFormat.SpaceBefore = 0
                                doc.Content.ParagraphFormat.SpaceAfter = 0
                                doc.Content.ParagraphFormat.LineSpacingRule = 0  # 0 = wdLineSpaceSingle
                                doc.Content.ParagraphFormat.WidowControl = True
                                doc.Content.ParagraphFormat.KeepTogether = False
                                doc.Content.ParagraphFormat.KeepWithNext = False
                                doc.Content.ParagraphFormat.PageBreakBefore = False
                            except Exception as e:
                                print(
                                    f"  └─ Ostrzeżenie: Błąd podczas zmiany formatowania ({e})"
                                )
                        print(f"  └─ Zakończono: {target.parent.name}/{target.name}")
                    finally:
                        doc.Save()
                        doc.Close(SaveChanges=False)
            # --- NAKŁADANIE MARGINESÓW NA SAMYM KOŃCU (PO WSZYSTKICH MAKRACH) ---
            print(">>> Aplikowanie ostatecznych marginesów z konfiguracji...")
            for f in files:
                f_upper = f.stem.upper()
                if f_upper in margin_config:
                    ext = ".docx" if f_upper == "REJESTR1" else ".doc"
                    rel_path = f.relative_to(in_dir)
                    target = out_dir / rel_path.parent / f"{f.stem}{ext}"
                    if target.exists():
                        doc = None
                        try:
                            doc = word.Documents.Open(str(target))
                            m = margin_config[f_upper]
                            doc.PageSetup.TopMargin = 28.35 * m[0]
                            doc.PageSetup.BottomMargin = 28.35 * m[1]
                            doc.PageSetup.LeftMargin = 28.35 * m[2]
                            doc.PageSetup.RightMargin = 28.35 * m[3]
                            doc.Save()
                            print(f"  └─ Ustawiono marginesy dla: {target.name}")
                        except Exception as e:
                            print(f"  └─ Ostrzeżenie: Błąd ustawiania marginesów ({target.name}): {e}")
                        finally:
                            if doc is not None:
                                doc.Close(SaveChanges=False)
            # --------------------------------------------------------------------
    finally:
        if word is not None:  # <--- Dodany warunek
            word.Quit()
        time.sleep(2)
        print(">>> Zakończono procesy tła Word.")

    print(">>> Optymalizacja układu czcionki (Autokorekta GUI)...")
    all_files = []
    for f in files:
        ext = ".docx" if f.stem.upper() == "REJESTR1" else ".doc"
        rel_path = f.relative_to(in_dir)
        target = out_dir / rel_path.parent / f"{f.stem}{ext}"
        if target.exists() and f.stem.upper() in FILES_TO_FIX:
            all_files.append(target)

    for idx, target in enumerate(all_files):
        os.startfile(str(target))

        # ZABEZPIECZENIE 1: Dłuższy czas na otworzenie pierwszego pliku
        if idx == 0:
            time.sleep(8)
        else:
            time.sleep(3.5)

        # ZABEZPIECZENIE 2: 'interval=0.15' zmusza skrypt do wyraźnego wciskania klawiszy
        pyautogui.hotkey("ctrl", "a", interval=0.15)
        time.sleep(0.3)
        pyautogui.press("alt")
        time.sleep(0.3)
        pyautogui.press("down")
        time.sleep(0.2)
        pyautogui.press("right")
        time.sleep(0.2)
        pyautogui.press("right")
        time.sleep(0.2)
        pyautogui.press("up")
        time.sleep(0.3)
        pyautogui.write("Cascadia Code")
        time.sleep(0.2)
        pyautogui.press("enter")
        time.sleep(0.5)

        # Odznaczamy tekst strzałką w prawo
        pyautogui.press("right")
        time.sleep(0.3)

        # ZABEZPIECZENIE 3: Wyraźne, powolne skróty z przerwami na zapis dyskowy
        pyautogui.hotkey("ctrl", "s", interval=0.2)
        time.sleep(1.2)  # Dajemy dyskowi ponad sekundę na zapisanie pliku
        pyautogui.hotkey("ctrl", "w", interval=0.2)
        time.sleep(0.8)

        print(f"  └─ Skorygowano wizualnie: {target.parent.name}/{target.name}")

    if all_files:
        time.sleep(1)
        try:
            word_app = win32com.client.GetActiveObject("Word.Application")
            word_app.Quit()
        except Exception:
            pyautogui.hotkey("alt", "f4", interval=0.2)
        print("  └─ Zakończono i zamknięto okno Microsoft Word.")




def get_resource_path(filename):
    candidates = []
    meipass = getattr(sys, "_MEIPASS", None)
    if meipass:
        candidates.append(Path(meipass) / filename)
    # Katalog główny projektu (dwa poziomy wyżej od app/core/word_worker.py)
    project_root = Path(__file__).resolve().parent.parent.parent
    candidates.append(project_root / filename)
    # Katalog, w którym znajduje się word_worker.py (dla kompatybilności)
    candidates.append(Path(__file__).resolve().parent / filename)
    candidates.append(Path.cwd() / filename)
    for path in candidates:
        if path.exists():
            return path
    return candidates[0]


def replace_text_preserve_runs(paragraph, replacements):
    full_text = "".join(run.text for run in paragraph.runs)
    if not full_text:
        return
    new_text = full_text
    for old, new in replacements.items():
        pattern = re.compile(re.escape(old), re.IGNORECASE)
        new_text = pattern.sub(new, new_text)
    if new_text == full_text:
        return
    if paragraph.runs:
        paragraph.runs[0].text = new_text
        for run in paragraph.runs[1:]:
            run.text = ""


def replace_text_in_tables(tables, replacements):
    for table in tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    replace_text_preserve_runs(paragraph, replacements)
                replace_text_in_tables(cell.tables, replacements)


